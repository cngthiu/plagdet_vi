# ============================================
# File: src/plagdet_vi/io/extract_docx.py
# ============================================
from __future__ import annotations
from docx import Document
from ..preprocess.normalize import normalize_text

def read_docx_text(path: str) -> str:
    d = Document(path)
    parts = [p.text for p in d.paragraphs]
    for table in d.tables:
        for row in table.rows:
            for cell in row.cells:
                val = cell.text.strip()
                if val:
                    parts.append(val)
    return normalize_text("\n".join(parts))